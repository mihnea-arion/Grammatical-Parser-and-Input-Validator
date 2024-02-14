import time
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from odf import text, teletype
from odf.opendocument import load

doc = docx.Document()

style = doc.styles['No Spacing']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

starters_integers = list(range(1,62))
starters = [' '+str(x)+') ' for x in starters_integers]

delimiters_normal = ['ORIZONTAL:', 'VERTICAL:', 'ORIZONTAL (', 'VERTICAL (']
delimiters_alea = ['De la dreapta la stânga: ', 'De la stânga la dreapta: ', 'De sus în jos']
title_proverbe = 'Proverbe ascunse'
title_cine_reconstituire = 'Cine-reconstituire'

detected_proverbe_ascunse = False
detected_cine_reconstituire = False

def replace_dash (paragraph):
        if (paragraph.text.find(' - ') != -1):
                moment = paragraph.text.replace(' - ', ' – ')
                paragraph.text = moment

def replace_start_dash (paragraph):
        if (paragraph.text.find('- ') != -1):
                moment = paragraph.text.replace('- ', '– ')
                paragraph.text = moment
        else:
                moment = '\t'+paragraph.text
                paragraph.text = moment

def place_tab (paragraph):
        if (paragraph.text.find('Exemplu: ') != -1):
                moment = paragraph.text.replace('Exemplu: ','')
                paragraph.text = '\tExemplu: '
                paragraph.runs[0].bold = True
                paragraph.add_run(moment)
        else:
                moment = '\t'+paragraph.text
                paragraph.text = moment
                
def easy_bold(paragraph, separator):
        splinter = paragraph.text.split(separator, 1)
        paragraph.text = splinter[0] + separator
        paragraph.runs[0].bold = True
        
        count = 2

        for x in starters:
                if x in splinter[1]:
                        paragraph.add_run(splinter[1].split(x, 1)[0])
                        splinter[1] = splinter[1].split(x, 1)[1]
                        paragraph.add_run(x)
                        paragraph.runs[count].bold = True
                        count = count + 2

        paragraph.add_run(splinter[1])

def italics_easy (paragraph):
        for ran in (paragraph.runs):
                ran.italic = True

def bold_titles_authors (paragraph):
        for ran in (paragraph.runs):
                ran.bold = True

def beginning_tab(paragraph):
        pass

def check_normal(paragraph):
        return (any(x in paragraph.text for x in delimiters_normal))

def check_dictionary (paragraph):
        return (paragraph.text.find('Dicționar: ') != -1)

def check_alea_iacta_est(paragraph):
        return (any(x in paragraph.text for x in delimiters_alea))

def check_spiral(paragraph):
        return(('1) ' in paragraph.text)and (' 40) ' in paragraph.text)
               and (check_normal(paragraph) == False) and (check_alea_iacta_est(paragraph) == False))

def check_proverbe_ascunse (paragraph, detected):
        if paragraph.text == title_proverbe:
                return True

def check_cine_reconstituire (paragraph, detected):
        if paragraph.text == title_cine_reconstituire:
                return True

odtdoc = load("C:/Users/Surubelnita/Desktop/input.odt")
allparas = odtdoc.getElementsByType(text.P)

for para in allparas:
        doc.add_paragraph(str(para))

for para in (doc.paragraphs):
        para.style = doc.styles['No Spacing']
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if (para.text!=''):
                replace_dash (para)

                #cine_reconstituire
                if detected_cine_reconstituire:
                        if para.text=='Mihai ZGUBEA':
                                detected_cine_reconstituire = False
                                bold_titles_authors (para)
                        else:
                                replace_start_dash (para)
                                pass
                #proverbe_ascunse
                elif detected_proverbe_ascunse:
                        if para.text=='Jocuri de Mihai ZGUBEA':
                                detected_proverbe_ascunse = False
                                bold_titles_authors (para)
                        else:
                                place_tab (para)
                                pass

                #horizontal_vertical
                elif check_normal(para):
                        easy_bold(para, ':')

                #dictionary
                elif check_dictionary(para):
                        italics_easy(para)

                #alea_iacta_est
                elif check_alea_iacta_est(para):
                        easy_bold(para, ':')

                #spirala
                elif check_spiral (para):
                        easy_bold(para, ' ')

                #check_proverbe_ascunse
                elif check_proverbe_ascunse (para, detected_proverbe_ascunse):
                        detected_proverbe_ascunse = check_proverbe_ascunse(para, detected_proverbe_ascunse)
                        bold_titles_authors (para)

                #check_cine_reconstituire
                elif check_cine_reconstituire (para, detected_cine_reconstituire):
                        detected_cine_reconstituire = check_cine_reconstituire(para, detected_cine_reconstituire)
                        bold_titles_authors (para)

                else:
                        bold_titles_authors (para)


        para.style = doc.styles['No Spacing']


doc.save('C:/Users/Surubelnita/Desktop/output.docx')


time.sleep(5)
