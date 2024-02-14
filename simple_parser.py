import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time

doc = docx.Document()

style = doc.styles['No Spacing']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

page_nos = []
values = {}

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

with open('C:/Users/Surubelnita/Desktop/input.txt', encoding='utf-8') as f:
    for line in f:
        doc.add_paragraph(line.strip('\n'))
    for para in doc.paragraphs:
        para.style = doc.styles['No Spacing']
        
    doc.save('C:/Users/Surubelnita/Desktop/output.docx')

doc = docx.Document('C:/Users/Surubelnita/Desktop/output.docx')

for line in (doc.paragraphs):
    r_line_split = line.text.split(' ', 1)

    page_nos.append(float(r_line_split[0]))
    if float(r_line_split[0]) != int(float(r_line_split[0])):
        values.update({float(r_line_split[0]) : '_DELIM_' + r_line_split[1]})
    else:
        values.update({float(r_line_split[0]) : 'Pag. ' + r_line_split[0] + '. ' + r_line_split[1]})

sorted_page_nos = sorted(page_nos)
sorted_values = dict(sorted(values.items()))

for page in sorted_values:
        if (float(page) != int(page)):
                sorted_values[int(page)] = sorted_values[int(page)] + ' ' + sorted_values[(page)]


i = 0
for line in (doc.paragraphs):
    if (float(sorted_page_nos[i]) == int(sorted_page_nos[i])):
        occurences = sorted_values[sorted_page_nos[i]].count('_DELIM_')

        scan = sorted_values[sorted_page_nos[i]].split(':', 1)
        line.text = scan[0]+':'
        line.runs[0].bold = True

        if (occurences == 0):
            line.add_run(scan[1])

        elif (occurences == 1):
            splinter = scan[1].split('_DELIM_', 1)
            line.add_run(splinter[0])

            second = splinter[1].split(':', 1)
            line.add_run(second[0]+':')
            line.runs[2].bold = True
            line.add_run(second[1])
 
        else:
            splinter = scan[1].split('_DELIM_', 2)
            line.add_run(splinter[0])

            second = splinter[1].split(':', 1)
            line.add_run(second[0]+':')
            line.runs[2].bold = True
            line.add_run(second[1])

            third = splinter[2].split(':', 1)
            line.add_run(third[0]+':')
            line.runs[4].bold = True
            line.add_run(third[1])

        line.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    else:
        delete_paragraph(line)
    i = i + 1

doc.save('C:/Users/Surubelnita/Desktop/output.docx')

time.sleep(5)
